import docx

doc = docx.Document('Контракт, перевод_рев_2_190319_1C.docx')
# print(len(doc.paragraphs))
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].text)
# print(doc.paragraphs[1].runs[0].text)
text = []
for paragraph in doc.paragraphs:
    if paragraph.text:
        for run in paragraph.runs:
            if run.text:
                text.append(paragraph.text)
print('\n\n'.join(text))
styles = []
for paragraph in doc.paragraphs:
    styles.append(paragraph.style)
print(*styles)